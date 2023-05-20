from docx import Document
from docx.shared import Inches
document = Document()

document.add_heading('簡単なWordドキュメントのタイトルasdasdsads', 0)
document.add_paragraph('簡単なWordドキュメントのテキスト')

document.save('sample.docx')
#画像
document.add_picture('_wsl.jpg', width=Inches(1.25))
#別名で保存
document.save('sample_answer.docx')


num = 0
for p in document.paragraphs:
    num += len(p.text)
print('文字数:',num)