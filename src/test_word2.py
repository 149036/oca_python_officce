import docx 
from docx.shared import Inches

#読み込み
document = docx.Document("sample.docx")
#画像
document.add_picture('_wsl.jpg', width=Inches(1.25))
#別名で保存
document.save('sample_answer.docx')

#カウント
num = 0
for p in document.paragraphs:
    num += len(p.text)
print('文字数:',num)